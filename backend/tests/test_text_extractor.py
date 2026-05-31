from pathlib import Path
import sys
import types

import pytest
from docx import Document

from app.config import settings
from app.services import ocr_service, text_extractor
from app.services.document_processor import process_document
from app.services.provider_errors import ProviderError
from app.services.text_extractor import (
    VLM_OCR_INFO_MESSAGE,
    extract_pages,
    extract_pages_with_metadata,
    extract_text,
)


class _FakePdfPage:
    def __init__(self, text: str) -> None:
        self._text = text

    def get_text(self) -> str:
        return self._text


class _FakePdfDocument:
    def __init__(self, page_texts: list[str]) -> None:
        self._pages = [_FakePdfPage(text) for text in page_texts]

    def __iter__(self):
        return iter(self._pages)

    def __len__(self) -> int:
        return len(self._pages)

    def __enter__(self):
        return self

    def __exit__(self, _exc_type, _exc, _tb) -> bool:
        return False


@pytest.fixture(autouse=True)
def reset_ocr_settings():
    snapshot = {
        "ocr_provider": settings.ocr_provider,
        "ocr_use_vlm": settings.ocr_use_vlm,
        "ocr_debug": settings.ocr_debug,
        "openrouter_api_key": settings.openrouter_api_key,
        "openrouter_model_ocr_vlm": settings.openrouter_model_ocr_vlm,
        "openrouter_ocr_model": settings.openrouter_ocr_model,
        "vision_provider": settings.vision_provider,
        "bothub_api_key": settings.bothub_api_key,
        "bothub_api_base_url": settings.bothub_api_base_url,
        "vision_api_base_url": settings.vision_api_base_url,
        "vision_api_key": settings.vision_api_key,
        "vision_model_ocr": settings.vision_model_ocr,
        "vision_timeout_seconds": settings.vision_timeout_seconds,
        "vision_include_usage": settings.vision_include_usage,
        "ocr_min_text_chars_per_page": settings.ocr_min_text_chars_per_page,
        "ocr_vlm_max_pages": settings.ocr_vlm_max_pages,
        "ocr_tesseract_lang": settings.ocr_tesseract_lang,
    }
    yield
    for key, value in snapshot.items():
        setattr(settings, key, value)


def _setup_scan_defaults(monkeypatch: pytest.MonkeyPatch, page_texts: list[str]) -> None:
    settings.ocr_provider = "hybrid"
    settings.ocr_use_vlm = True
    settings.vision_provider = "bothub"
    settings.bothub_api_base_url = "https://openai.bothub.chat/v1"
    settings.bothub_api_key = "dummy"
    settings.vision_api_base_url = "https://openai.bothub.chat/v1"
    settings.vision_api_key = "dummy"
    settings.vision_model_ocr = "gpt-4o"
    settings.ocr_vlm_max_pages = 20
    settings.ocr_min_text_chars_per_page = 50
    monkeypatch.setattr(text_extractor.fitz, "open", lambda _path: _FakePdfDocument(page_texts))
    monkeypatch.setattr(text_extractor, "_render_pdf_page_to_png_bytes", lambda _page, _dpi: b"png")


def test_extract_text_from_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("Payment terms are net 30.")
    document.add_paragraph("Termination requires 30 days notice.")
    document.save(str(file_path))

    extracted = extract_text(str(file_path))
    assert "Payment terms are net 30." in extracted
    assert "Termination requires 30 days notice." in extracted


def test_docx_extraction_creates_chunks(tmp_path: Path) -> None:
    file_path = tmp_path / "large_contract.docx"
    document = Document()
    for index in range(120):
        document.add_paragraph(
            f"Clause {index}: Payment must be made within 10 banking days and liability applies for delay."
        )
    document.save(str(file_path))

    processed = process_document("docx-large-1", str(file_path))

    assert processed["text_length"] > 1000
    assert processed["chunks_count"] > 0
    assert processed["chunk_records"]


def test_pdf_scan_uses_vlm_ocr_when_text_layer_empty(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    monkeypatch.setattr(
        text_extractor,
        "extract_text_from_image_with_vision_provider",
        lambda **_kwargs: "Стороны согласовали срок оплаты в течение десяти банковских дней и условия ответственности.",
    )
    monkeypatch.setattr(
        text_extractor,
        "run_ocr_image_bytes",
        lambda _image_bytes: (_ for _ in ()).throw(AssertionError("Tesseract should not be used")),
    )

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-1")
    assert result["used_ocr"] is True
    assert result["pages"][0]["source"] == "vlm_ocr"
    assert "срок оплаты" in result["pages"][0]["text"].lower()


def test_vlm_ocr_text_is_accepted_and_chunked(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    vlm_text = (
        "Договор поставки. Поставщик обязуется поставить товар. Покупатель обязуется оплатить товар "
        "в течение 10 банковских дней с даты счета. Ответственность сторон определяется условиями договора."
    )
    monkeypatch.setattr(text_extractor, "extract_text_from_image_with_vision_provider", lambda **_kwargs: vlm_text)
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "")

    processed = process_document("doc-ocr-2", "scan.pdf")
    assert processed["text_length"] > 80
    assert processed["chunks_count"] > 0
    assert processed["pages"][0]["source"] == "vlm_ocr"


def test_vlm_ocr_meta_response_is_rejected_and_tesseract_used(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    monkeypatch.setattr(
        text_extractor,
        "extract_text_from_image_with_vision_provider",
        lambda **_kwargs: "На изображении видно договор с условиями оплаты и ответственностью сторон.",
    )
    monkeypatch.setattr(
        text_extractor,
        "run_ocr_image_bytes",
        lambda _image_bytes: "Оплата производится в течение 10 банковских дней.",
    )

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-3")
    assert result["pages"][0]["source"] == "tesseract"


def test_vlm_ocr_refusal_is_rejected_and_tesseract_used(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    monkeypatch.setattr(
        text_extractor,
        "extract_text_from_image_with_vision_provider",
        lambda **_kwargs: "I'm unable to read this image clearly.",
    )
    monkeypatch.setattr(
        text_extractor,
        "run_ocr_image_bytes",
        lambda _image_bytes: "Договор вступает в силу с момента подписания.",
    )

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-4")
    assert result["pages"][0]["source"] == "tesseract"


def test_vlm_ocr_mojibake_rejected(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    monkeypatch.setattr(text_extractor, "extract_text_from_image_with_vision_provider", lambda **_kwargs: "РџСЂРёРІРµС‚")
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "Срок оплаты составляет 5 дней.")

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-5")
    assert result["pages"][0]["source"] == "tesseract"
    assert any("Vision OCR" in warning for warning in result["warnings"])


def test_successful_vlm_ocr_is_not_warning_by_itself(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    monkeypatch.setattr(
        text_extractor,
        "extract_text_from_image_with_vision_provider",
        lambda **_kwargs: "Покупатель оплачивает поставку в срок, установленный договором и спецификацией.",
    )
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "")

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-6")
    non_info = [w for w in result["warnings"] if not str(w).startswith("INFO:")]
    assert non_info == []
    assert VLM_OCR_INFO_MESSAGE in result["warnings"]


def test_insufficient_ocr_text_warning_is_human_readable_ru(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    monkeypatch.setattr(text_extractor, "extract_text_from_image_with_vision_provider", lambda **_kwargs: "На изображении видно договор")
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "...")

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-7")
    assert any("Из PDF-скана удалось извлечь недостаточно текста" in warning for warning in result["warnings"])


def test_ocr_page_order_is_preserved(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, ["", ""])

    texts = {
        1: "Страница 1: Предмет договора.",
        2: "Страница 2: Ответственность сторон.",
    }

    monkeypatch.setattr(
        text_extractor,
        "extract_text_from_image_with_vision_provider",
        lambda **kwargs: texts[int(kwargs["page_number"])],
    )
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "")

    result = extract_pages_with_metadata("scan.pdf", document_id="doc-ocr-8")
    assert len(result["pages"]) == 2
    assert result["pages"][0]["text"].startswith("Страница 1")
    assert result["pages"][1]["text"].startswith("Страница 2")


def test_missing_vision_model_falls_back_to_tesseract_without_provider_call(monkeypatch: pytest.MonkeyPatch) -> None:
    settings.ocr_provider = "hybrid"
    settings.ocr_use_vlm = True
    settings.vision_provider = "bothub"
    settings.bothub_api_key = "dummy"
    settings.vision_api_key = "dummy"
    settings.vision_model_ocr = ""

    monkeypatch.setattr(text_extractor.fitz, "open", lambda _path: _FakePdfDocument([""]))
    monkeypatch.setattr(text_extractor, "_render_pdf_page_to_png_bytes", lambda _page, _dpi: b"png")
    monkeypatch.setattr(
        text_extractor,
        "extract_text_from_image_with_vision_provider",
        lambda **_kwargs: (_ for _ in ()).throw(AssertionError("Provider must not be called when model is missing")),
    )
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "Fallback tesseract text")

    result = extract_pages_with_metadata("scanned.pdf")
    assert result["pages"][0]["source"] == "tesseract"


def test_provider_rate_limited_falls_back_to_tesseract(monkeypatch: pytest.MonkeyPatch) -> None:
    _setup_scan_defaults(monkeypatch, [""])

    def _raise_429(**_kwargs):
        raise ProviderError(
            provider="bothub",
            code="provider_rate_limited",
            message="rate limited",
            retryable=True,
        )

    monkeypatch.setattr(text_extractor, "extract_text_from_image_with_vision_provider", _raise_429)
    monkeypatch.setattr(text_extractor, "run_ocr_image_bytes", lambda _image_bytes: "Fallback tesseract text")

    result = extract_pages_with_metadata("scanned.pdf")
    assert result["pages"][0]["source"] == "tesseract"


def test_extract_pages_from_docx_uses_single_page(tmp_path: Path) -> None:
    file_path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("Clause on page one.")
    document.save(str(file_path))

    pages = extract_pages(str(file_path))
    assert len(pages) == 1
    assert pages[0]["page"] == 1
    assert pages[0]["source"] == "docx"
    assert "Clause on page one." in pages[0]["text"]


def test_extract_text_unsupported_extension_raises_value_error(tmp_path: Path) -> None:
    file_path = tmp_path / "contract.txt"
    file_path.write_text("plain text", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(str(file_path))


def test_run_ocr_image_bytes_uses_default_rus_plus_eng_lang(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, str] = {}

    class _DummyImage:
        def __enter__(self):
            return object()

        def __exit__(self, _exc_type, _exc, _tb):
            return False

    class _PytesseractInner:
        tesseract_cmd = ""

    def _image_to_string(_img, lang=None):
        captured["lang"] = str(lang or "")
        return "ok"

    pytesseract_module = types.ModuleType("pytesseract")
    pytesseract_module.pytesseract = _PytesseractInner()
    pytesseract_module.image_to_string = _image_to_string

    pil_image_module = types.ModuleType("PIL.Image")
    pil_image_module.open = lambda _stream: _DummyImage()
    pil_module = types.ModuleType("PIL")
    pil_module.Image = pil_image_module

    monkeypatch.setitem(sys.modules, "pytesseract", pytesseract_module)
    monkeypatch.setitem(sys.modules, "PIL", pil_module)
    monkeypatch.setitem(sys.modules, "PIL.Image", pil_image_module)

    settings.ocr_tesseract_lang = "rus+eng"
    result = ocr_service.run_ocr_image_bytes(b"fake")

    assert captured["lang"] == "rus+eng"
    assert result == "ok"


def test_run_ocr_image_bytes_falls_back_to_eng_when_lang_pack_missing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    calls: list[str] = []

    class _DummyImage:
        def __enter__(self):
            return object()

        def __exit__(self, _exc_type, _exc, _tb):
            return False

    class _PytesseractInner:
        tesseract_cmd = ""

    def _image_to_string(_img, lang=None):
        calls.append(str(lang or ""))
        if lang == "rus+eng":
            raise RuntimeError("Failed loading language 'rus+eng'")
        return "english-fallback-ok"

    pytesseract_module = types.ModuleType("pytesseract")
    pytesseract_module.pytesseract = _PytesseractInner()
    pytesseract_module.image_to_string = _image_to_string

    pil_image_module = types.ModuleType("PIL.Image")
    pil_image_module.open = lambda _stream: _DummyImage()
    pil_module = types.ModuleType("PIL")
    pil_module.Image = pil_image_module

    monkeypatch.setitem(sys.modules, "pytesseract", pytesseract_module)
    monkeypatch.setitem(sys.modules, "PIL", pil_module)
    monkeypatch.setitem(sys.modules, "PIL.Image", pil_image_module)

    settings.ocr_tesseract_lang = "rus+eng"
    result = ocr_service.run_ocr_image_bytes(b"fake")

    assert calls == ["rus+eng", "eng"]
    assert result == "english-fallback-ok"
