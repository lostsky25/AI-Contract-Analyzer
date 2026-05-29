from pathlib import Path

from docx import Document


def write_demo_contract(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("ДОГОВОР ПОСТАВКИ № 1")
    document.add_paragraph(
        "Стороны договорились о поставке товара в течение 30 календарных дней."
    )
    document.add_paragraph(
        "Оплата производится в течение 10 банковских дней с даты поставки."
    )
    document.add_paragraph(
        "За просрочку оплаты предусмотрена неустойка 0.1% за каждый день просрочки."
    )
    document.add_paragraph(
        "Договор может быть расторгнут в одностороннем порядке при существенном нарушении."
    )
    document.add_paragraph(
        "Конфиденциальная информация не подлежит разглашению третьим лицам."
    )
    document.save(str(path))
    return path


def demo_contract_path() -> Path:
    return Path(__file__).resolve().parents[2] / "test_data" / "demo_contract.docx"
